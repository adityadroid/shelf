from __future__ import annotations

from pathlib import Path

from docx import Document

from shelf.indexing.models import DocumentSection, ParsedDocument, ParserStatus
from shelf.parsers.base import DocumentParser


class DocxParser(DocumentParser):
    parser_type = "python-docx"

    def parse(self, path: Path) -> ParsedDocument:
        stats = path.stat()
        diagnostics: list[str] = []
        text_parts: list[str] = []
        sections: list[DocumentSection] = []

        try:
            document = Document(str(path))
            for index, paragraph in enumerate(document.paragraphs, start=1):
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
                    sections.append(DocumentSection(source_ref=f"paragraph:{index}", text=text))

            for table_index, table in enumerate(document.tables, start=1):
                for row_index, row in enumerate(table.rows, start=1):
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
                        sections.append(
                            DocumentSection(source_ref=f"table:{table_index}:row:{row_index}", text=row_text)
                        )

            status = ParserStatus.SUCCESS if text_parts else ParserStatus.NO_TEXT
            if status == ParserStatus.NO_TEXT:
                diagnostics.append("No readable text found in DOCX.")
        except Exception as exc:
            diagnostics.append(str(exc))
            status = ParserStatus.FAILURE

        return ParsedDocument(
            path=str(path),
            file_name=path.name,
            extension=path.suffix.lower(),
            size_bytes=stats.st_size,
            ctime=stats.st_ctime,
            mtime=stats.st_mtime,
            parser_type=self.parser_type,
            parser_status=status,
            raw_text="\n\n".join(text_parts),
            diagnostics=diagnostics,
            sections=sections,
        )
